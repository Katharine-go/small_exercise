# 1.将wed/html内容导出为world文档
#   (1)unoconv:支持将本地html文档转换为docx格式的文档，所以需要先将网页中的html文件保存到本地，再调用unoconv进行转换。

"""
  安装：sudo apt-get install unoconv
  使用：unoconv -f pdf *.odt
       unoconv -f doc *.odt
       unoconv -f html *.odt
"""

# (2)python-docx :python-docx是一个可以读写word文档的python库。
#                 获取网页中的数据，使用python手动排版添加到word文档中。
from docx import Document
from docx.shared import Inches
document = Document()
document.add_heading('Document Title', 0)
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')
document.add_paragraph(
 'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
 'first item in ordered list', style='ListNumber'
)
document.add_picture('monty-truth.png', width=Inches(1.25))
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in recordset:
 row_cells = table.add_row().cells
 row_cells[0].text = str(item.qty)
 row_cells[1].text = str(item.id)
 row_cells[2].text = item.desc
document.add_page_break()
document.save('demo.docx')

from docx import Document
from docx.shared import Inches
document = Document()
for row in range(9):
 t = document.add_table(rows=1,cols=1,style = 'Table Grid')
 t.autofit = False #很重要！
 w = float(row) / 2.0
 t.columns[0].width = Inches(w)
document.save('table-step.docx')

# 2.导出PDF文档方法
# (1)pdfkit：
#   a.wkhtmltopdf主要用于HTML生成PDF。
#   b.pdfkit是基于wkhtmltopdf的python封装，支持URL，本地文件，文本内容到PDF的转换，其最终还是调用wkhtmltopdf命令。是目前接触到的python生成pdf效果较好的。

"""
  安装 pip install pdfkit
  使用：
      import pdfkit
      pdfkit.from_url('http://google.com', 'out.pdf')
      pdfkit.from_file('test.html', 'out.pdf')
      pdfkit.from_string('Hello!', 'out.pdf')
"""